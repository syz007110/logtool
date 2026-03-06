import argparse
import base64
import hashlib
import hmac
import json
import os
import random
import re
import sqlite3
import sys
import time
import urllib.request
import urllib.error
import urllib.parse
import xml.etree.ElementTree as ET
from concurrent.futures import ThreadPoolExecutor
from threading import Lock

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


PROMPT_VERSION = "v2"
LOG_TIMING = str(os.environ.get("TRANSLATE_LOG_TIMING", "")).strip().lower() in ("1", "true", "yes", "on")
LINE_MERGE_DELIM = "<<<LN>>>"  # 非 LLM 逐行合并时使用的行分隔符，翻译后需还原为 \n

# 讯飞等 MT 可能将 <<<LN>>> 改写为 <LN>、<<、< < 等，需容错还原
_LINE_DELIM_PATTERN = re.compile(
    r"<<<\s*LN\s*>>>|<\s*LN\s*>|<\s*<",
    re.IGNORECASE,
)


def _restore_line_merge_delim(text: str) -> str:
    """将 MT 可能改写后的行分隔符还原为 \\n"""
    out = text.replace(LINE_MERGE_DELIM, "\n")
    out = _LINE_DELIM_PATTERN.sub("\n", out)
    return out


def _norm(s: str) -> str:
    return str(s or "")


def sha256_text(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8")).hexdigest()


def normalize_text_whitespace(text: str) -> str:
    """
    Light normalization for cache hit improvements:
    - strip trailing spaces per line
    - collapse consecutive spaces/tabs inside lines
    - preserve newlines
    """
    lines = text.splitlines(keepends=True)
    out_lines: List[str] = []
    for line in lines:
        if line.endswith("\r\n"):
            core = line[:-2]
            line_end = "\r\n"
        elif line.endswith("\n"):
            core = line[:-1]
            line_end = "\n"
        elif line.endswith("\r"):
            core = line[:-1]
            line_end = "\r"
        else:
            core = line
            line_end = ""
        core = core.rstrip()
        core = re.sub(r"[ \t]+", " ", core)
        out_lines.append(core + line_end)
    return "".join(out_lines)


def safe_json_load(path: Path) -> Any:
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None


def ensure_dir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)


def normalize_base_url(base_url: str) -> str:
    s = str(base_url or "").strip()
    return s.rstrip("/")


def clamp_int(n: Any, min_v: int, max_v: int, fallback: int) -> int:
    try:
        v = int(str(n))
    except Exception:
        return fallback
    return max(min(v, max_v), min_v)


@dataclass
class Provider:
    id: str
    label: str
    kind: str
    requires_api_key: bool
    api_key: str
    app_id: str
    api_secret: str
    base_url: str
    model: str
    timeout_ms: int
    temperature: float
    top_p: float


def load_providers(providers_file: Path) -> List[Provider]:
    raw = safe_json_load(providers_file)
    if not isinstance(raw, list):
        raise RuntimeError(f"providers-file invalid or not found: {providers_file}")

    out: List[Provider] = []
    seen = set()
    for spec in raw:
        if not isinstance(spec, dict):
            continue
        pid = str(spec.get("id") or "").strip()
        if not pid:
            continue
        key = pid.lower()
        if key in seen:
            continue
        seen.add(key)

        api_key_env = str(spec.get("apiKeyEnv") or "").strip()
        api_key = str(spec.get("apiKey") or (os.environ.get(api_key_env, "") if api_key_env else "") or "").strip()
        app_id_env = str(spec.get("appIdEnv") or "").strip()
        app_id = str(spec.get("appId") or (os.environ.get(app_id_env, "") if app_id_env else "") or "").strip()
        api_secret_env = str(spec.get("apiSecretEnv") or "").strip()
        api_secret = str(spec.get("apiSecret") or (os.environ.get(api_secret_env, "") if api_secret_env else "") or "").strip()
        base_url = normalize_base_url(str(spec.get("baseUrl") or "").strip())
        model = str(spec.get("model") or "").strip()
        label = str(spec.get("label") or pid).strip() or pid
        kind = str(spec.get("kind") or "openai_compatible").strip()
        requires_api_key = bool(spec.get("requiresApiKey", True))
        timeout_ms = clamp_int(spec.get("timeoutMs"), 1000, 120000, 12000)
        temperature = float(spec.get("temperature", 0) or 0)
        top_p = float(spec.get("top_p", spec.get("topP", 0.1)) or 0.1)

        out.append(
            Provider(
                id=pid,
                label=label,
                kind=kind,
                requires_api_key=requires_api_key,
                api_key=api_key,
                app_id=app_id,
                api_secret=api_secret,
                base_url=base_url,
                model=model,
                timeout_ms=timeout_ms,
                temperature=temperature,
                top_p=top_p,
            )
        )
    if not out:
        raise RuntimeError(f"providers-file empty after filtering: {providers_file}")
    return out


def resolve_provider(providers: List[Provider], want_id: Optional[str]) -> Provider:
    want = str(want_id or "").strip()
    default_id = str(os.environ.get("SMART_SEARCH_LLM_DEFAULT_PROVIDER", "") or "").strip()

    if want:
        for p in providers:
            if p.id == want:
                return p
    if default_id:
        for p in providers:
            if p.id == default_id:
                return p
    return providers[0]


def llm_available(p: Provider) -> Tuple[bool, str]:
    dry_run = str(os.environ.get("TRANSLATE_DRY_RUN", "")).strip().lower() in ("1", "true", "yes", "on")
    if dry_run:
        return True, "dry_run"
    enabled = str(os.environ.get("SMART_SEARCH_LLM_ENABLED", "true")).strip().lower() not in ("0", "false", "no", "off")
    if not enabled:
        return False, "not_enabled"
    if p.kind == "xfyun":
        if not p.app_id:
            return False, "missing_app_id"
        if not p.api_key:
            return False, "missing_api_key"
        if not p.api_secret:
            return False, "missing_api_secret"
        if not p.base_url:
            return False, "missing_base_url"
        return True, "ok"
    if p.requires_api_key and not p.api_key:
        return False, "missing_api_key"
    if not p.base_url:
        return False, "missing_base_url"
    if not p.model:
        return False, "missing_model"
    return True, "ok"


def _log_timing(label: str, start_ts: float) -> None:
    if not LOG_TIMING:
        return
    try:
        ms = int((time.perf_counter() - start_ts) * 1000)
        sys.stderr.write(f"[translate-kernel] {label}: {ms}ms\n")
    except Exception:
        pass


class SqliteCache:
    def __init__(self, db_path: Path):
        ensure_dir(db_path.parent)
        self.db_path = db_path
        self.conn = sqlite3.connect(str(db_path), check_same_thread=False)
        self.lock = Lock()
        self._init()

    def _init(self) -> None:
        cur = self.conn.cursor()
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS translation_cache (
              cache_key TEXT PRIMARY KEY,
              created_at INTEGER NOT NULL,
              provider_id TEXT NOT NULL,
              model TEXT NOT NULL,
              lang_in TEXT NOT NULL,
              lang_out TEXT NOT NULL,
              glossary_hash TEXT NOT NULL,
              prompt_version TEXT NOT NULL,
              original_text TEXT NOT NULL,
              translated_text TEXT NOT NULL
            )
            """
        )
        self.conn.commit()

    def get(self, cache_key: str) -> Optional[str]:
        with self.lock:
            cur = self.conn.cursor()
            cur.execute("SELECT translated_text FROM translation_cache WHERE cache_key = ?", (cache_key,))
            row = cur.fetchone()
            return row[0] if row else None

    def set(self, cache_key: str, payload: Dict[str, Any]) -> None:
        with self.lock:
            cur = self.conn.cursor()
            cur.execute(
                """
                INSERT OR REPLACE INTO translation_cache
                (cache_key, created_at, provider_id, model, lang_in, lang_out, glossary_hash, prompt_version, original_text, translated_text)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    cache_key,
                    int(time.time()),
                    payload["provider_id"],
                    payload["model"],
                    payload["lang_in"],
                    payload["lang_out"],
                    payload["glossary_hash"],
                    payload["prompt_version"],
                    payload["original_text"],
                    payload["translated_text"],
                ),
            )
            self.conn.commit()


def glossary_load_from_sqlite(
    glossary_db: Path,
    source_lang: str,
    target_lang: str,
    domain: Optional[str] = None,
) -> Tuple[List[Tuple[str, str]], List[Tuple[str, List[str]]], str, Dict[str, Any]]:
    """
    SQLite glossary loader (compatible with randomTranslate schema):
    - term_concept(id, domain, ...)
    - term_lexeme(concept_id, lang, text, priority, status, ...)
    """
    entries: List[Tuple[str, str]] = []
    synonyms: List[Tuple[str, List[str]]] = []
    options: Dict[str, Any] = {}

    if not glossary_db.exists():
        norm_obj = {
            "mode": "sqlite",
            "db": str(glossary_db),
            "source_lang": source_lang,
            "target_lang": target_lang,
            "domain": domain or "",
            "entries": [],
        }
        glossary_hash = sha256_text(json.dumps(norm_obj, ensure_ascii=False, separators=(",", ":")))
        return entries, synonyms, glossary_hash, options

    src = normalize_lang(source_lang)
    tgt = normalize_lang(target_lang)
    conn = sqlite3.connect(str(glossary_db))
    try:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT ls.text AS src_text, lt.text AS tgt_text
            FROM term_concept c
            JOIN term_lexeme ls ON ls.concept_id = c.id AND ls.lang = ? AND ls.status = 'approved'
            JOIN term_lexeme lt ON lt.concept_id = c.id AND lt.lang = ? AND lt.status = 'approved'
            WHERE (? IS NULL OR c.domain = ?)
            ORDER BY LENGTH(ls.text) DESC, ls.priority ASC
            """,
            (src, tgt, domain, domain),
        )
        rows = cur.fetchall()
        for row in rows:
            s = str((row[0] if len(row) > 0 else "") or "").strip()
            t = str((row[1] if len(row) > 1 else "") or "").strip()
            if s and t:
                entries.append((s, t))
    finally:
        conn.close()

    norm_obj = {
        "mode": "sqlite",
        "db": str(glossary_db),
        "source_lang": src,
        "target_lang": tgt,
        "domain": domain or "",
        "entries": sorted(entries, key=lambda x: (x[0], x[1])),
    }
    glossary_hash = sha256_text(json.dumps(norm_obj, ensure_ascii=False, separators=(",", ":")))
    return entries, synonyms, glossary_hash, options


def apply_synonym_fixes(text: str, synonyms: List[Tuple[str, List[str]]]) -> str:
    out = text
    for canonical, variants in synonyms:
        for v in variants:
            if not v:
                continue
            # keep simple: case-sensitive for now; users can add variants
            out = out.replace(v, canonical)
    return out


def _should_use_word_boundary(term: str) -> bool:
    # Use word boundaries only for ascii-ish terms to avoid breaking CJK matches.
    return any(ch.isalnum() for ch in term) and all(ord(ch) < 128 for ch in term)


def make_term_placeholders(
    text: str,
    entries: List[Tuple[str, str]],
    options: Optional[Dict[str, Any]] = None,
) -> Tuple[str, Dict[str, str], List[str]]:
    """
    Replace matched source terms with placeholders {{T0001}}.
    Returns: (text_with_placeholders, placeholder_to_target, used_sources)
    """
    if not entries:
        return text, {}, []

    # longest first to avoid partial overlap
    sorted_entries = sorted(entries, key=lambda x: len(x[0]), reverse=True)
    placeholder_map: Dict[str, str] = {}
    used_sources: List[str] = []
    out = text
    idx = 1
    use_ci = bool((options or {}).get("caseInsensitive", False))
    use_boundary = bool((options or {}).get("wordBoundary", False))
    for src, tgt in sorted_entries:
        if not src:
            continue
        ph = f"{{{{T{idx:04d}}}}}"
        idx += 1
        if use_ci or use_boundary:
            flags = re.IGNORECASE if use_ci else 0
            pattern = re.escape(src)
            if use_boundary and _should_use_word_boundary(src):
                pattern = r"\b" + pattern + r"\b"
            new_out, n = re.subn(pattern, ph, out, flags=flags)
            if n > 0:
                out = new_out
                placeholder_map[ph] = tgt
                used_sources.append(src)
        else:
            if src in out:
                out = out.replace(src, ph)
                placeholder_map[ph] = tgt
                used_sources.append(src)
    return out, placeholder_map, used_sources


def restore_placeholders(text: str, placeholder_map: Dict[str, str]) -> str:
    """Restore terminology placeholders. Handles Xunfei/MT output variants."""
    out = text
    for ph, tgt in placeholder_map.items():
        # 1) strict restore first for exact placeholders, e.g. {{T0001}}
        out = out.replace(ph, tgt)
        m = re.match(r"\{\{T(\d{4})\}\}", ph)
        if not m:
            continue
        idx = m.group(1)
        patterns = [
            # "{{ T0001 }}" (double braces, optional spaces)
            re.compile(r"\{\{\s*T\s*" + idx + r"\s*\}\}", re.IGNORECASE),
            # "{ { T0001 } }" (space between braces - Xunfei output)
            re.compile(r"\{\s*\{\s*T\s*" + idx + r"\s*\}\s*\}", re.IGNORECASE),
            # "{ T0001 }" (single braces - Xunfei output)
            re.compile(r"\{\s*T\s*" + idx + r"\s*\}", re.IGNORECASE),
            # "\{\{T0001\}\}" (escaped)
            re.compile(r"\\\{\s*\\\{\s*T\s*" + idx + r"\s*\\\}\s*\\\}", re.IGNORECASE),
        ]
        for pattern in patterns:
            out = pattern.sub(tgt, out)
    return out


def protect_markdown(text: str) -> Tuple[str, Dict[str, str]]:
    """
    Protect code blocks, inline code, and link urls.
    Returns: (protected_text, token_map)
    """
    token_map: Dict[str, str] = {}
    counter = 0

    def token(prefix: str) -> str:
        nonlocal counter
        counter += 1
        return f"[[{prefix}_{counter:04d}]]"

    # fenced code blocks ```...```
    def repl_fenced(m: re.Match) -> str:
        t = token("CODEBLOCK")
        token_map[t] = m.group(0)
        return t

    out = re.sub(r"```[\s\S]*?```", repl_fenced, text)

    # HTML pre/code blocks
    def repl_pre(m: re.Match) -> str:
        t = token("HTMLBLOCK")
        token_map[t] = m.group(0)
        return t

    out = re.sub(r"<pre[\s\S]*?</pre>", repl_pre, out, flags=re.IGNORECASE)
    out = re.sub(r"<code[\s\S]*?</code>", repl_pre, out, flags=re.IGNORECASE)

    # inline code `...` (single-line)
    def repl_inline(m: re.Match) -> str:
        t = token("INLINE")
        token_map[t] = m.group(0)
        return t

    out = re.sub(r"`[^`\n]+`", repl_inline, out)

    # protect markdown link urls: ](url)
    def repl_url(m: re.Match) -> str:
        url = m.group(1)
        t = token("URL")
        token_map[t] = url
        return f"]({t})"

    out = re.sub(r"\]\(([^)]+)\)", repl_url, out)
    return out, token_map


def restore_markdown(text: str, token_map: Dict[str, str]) -> str:
    out = text
    # restore urls first (they're inside parentheses)
    for k, v in token_map.items():
        out = out.replace(k, v)
    # tolerant restore for model-altered markdown protection tokens, e.g.
    # "[[ URL_0001 ]]", "[ [INLINE_0002] ]", "[[codeblock 0003]]"
    for k, v in token_map.items():
        m = re.match(r"\[\[([A-Z]+)_(\d{4})\]\]", k)
        if not m:
            continue
        prefix, idx = m.group(1), m.group(2)
        pattern = re.compile(
            rf"\[\s*\[\s*{re.escape(prefix)}\s*[_\-\s]*{idx}\s*\]\s*\]",
            re.IGNORECASE,
        )
        out = pattern.sub(v, out)
    return out


def split_paragraphs_keep_separators(text: str) -> List[Tuple[str, bool]]:
    """
    Returns list of (segment, is_separator).
    Separator segments are blank lines (\\n\\n...), preserved as-is.
    """
    parts: List[Tuple[str, bool]] = []
    pattern = re.compile(r"(\n\s*\n+)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts


def merge_short_lines(
    lines: List[str],
    min_chars: int,
    max_chars: int,
) -> List[Tuple[str, bool]]:
    """
    合并短行用于翻译，用 LINE_MERGE_DELIM 连接，翻译后需 split 还原为 \\n。
    空行不参与合并，作为自然边界。
    """
    if min_chars <= 0:
        segments = []
        for i, line in enumerate(lines):
            if i > 0:
                segments.append(("\n", True))
            segments.append((line, False))
        return segments
    out: List[Tuple[str, bool]] = []
    buf: List[str] = []
    buf_len = 0

    def flush_buf() -> None:
        nonlocal buf, buf_len
        if buf:
            merged = LINE_MERGE_DELIM.join(buf)
            out.append((merged, False))
            buf, buf_len = [], 0

    for i, line in enumerate(lines):
        line_len = len(line)
        is_empty = not line.strip()
        if is_empty:
            flush_buf()
            out.append((line, False))
            if i < len(lines) - 1:
                out.append(("\n", True))
            continue
        if not buf:
            if line_len < min_chars:
                buf = [line]
                buf_len = line_len
            else:
                out.append((line, False))
                if i < len(lines) - 1:
                    out.append(("\n", True))
            continue
        add_len = 1 + line_len
        if buf_len + add_len <= max_chars and (line_len < min_chars or buf_len < min_chars):
            buf.append(line)
            buf_len += add_len
        else:
            flush_buf()
            out.append(("\n", True))
            if line_len < min_chars:
                buf = [line]
                buf_len = line_len
            else:
                out.append((line, False))
                if i < len(lines) - 1:
                    out.append(("\n", True))
    flush_buf()
    return out


def merge_short_segments(
    parts: List[Tuple[str, bool]],
    min_chars: int,
    max_chars: int,
) -> List[Tuple[str, bool]]:
    if min_chars <= 0:
        return parts
    out: List[Tuple[str, bool]] = []
    buf = ""
    for seg, is_sep in parts:
        if is_sep:
            if buf:
                buf += seg
            else:
                out.append((seg, True))
            continue
        if not seg:
            if buf:
                buf += seg
            else:
                out.append((seg, False))
            continue

        seg_len = len(seg)
        if not buf:
            if seg_len < min_chars:
                buf = seg
            else:
                out.append((seg, False))
            continue

        if len(buf) + seg_len <= max_chars and (seg_len < min_chars or len(buf) < min_chars):
            buf += seg
        else:
            out.append((buf, False))
            if seg_len < min_chars:
                buf = seg
            else:
                out.append((seg, False))
                buf = ""

    if buf:
        out.append((buf, False))
    return out


def split_long_text(text: str, max_chars: int) -> List[str]:
    s = text
    if len(s) <= max_chars:
        return [s]

    # split by sentence punctuation
    seps = re.compile(r"([???!?]+[\s\u3000]*|\n)")
    chunks: List[str] = []
    buf = ""
    for part in seps.split(s):
        if not part:
            continue
        if len(buf) + len(part) <= max_chars:
            buf += part
        else:
            if buf:
                chunks.append(buf)
                buf = ""
            if len(part) <= max_chars:
                buf = part
            else:
                # hard split
                for i in range(0, len(part), max_chars):
                    chunks.append(part[i : i + max_chars])
                buf = ""
    if buf:
        chunks.append(buf)
    return chunks


def _default_system_prompt(source_lang: str, target_lang: str) -> str:
    return (
        "You are a professional translation engine.\n"
        f"Translate from {source_lang} to {target_lang}.\n"
        "Rules:\n"
        "- Output ONLY the translated text. No explanations.\n"
        "- Preserve placeholders like {{T0001}} exactly.\n"
        "- Preserve tokens like [[CODEBLOCK_0001]] and [[INLINE_0001]] and [[URL_0001]] exactly.\n"
        "- Keep original formatting as much as possible.\n"
    )


def build_translation_messages(
    source_lang: str,
    target_lang: str,
    text: str,
    system_prompt: Optional[str] = None,
) -> List[Dict[str, str]]:
    # Keep prompt minimal to save tokens.
    if not system_prompt:
        system_prompt = _default_system_prompt(source_lang, target_lang)
    system = system_prompt
    return [
        {"role": "system", "content": system},
        {"role": "user", "content": text},
    ]


def _retry_params() -> Tuple[int, int, float]:
    max_retries = clamp_int(os.environ.get("TRANSLATE_RETRY_MAX"), 0, 6, 2)
    backoff_ms = clamp_int(os.environ.get("TRANSLATE_RETRY_BACKOFF_MS"), 100, 10000, 500)
    backoff_mult = float(os.environ.get("TRANSLATE_RETRY_BACKOFF_MULT", "2.0") or 2.0)
    return max_retries, backoff_ms, backoff_mult


def _sleep_backoff(attempt: int, base_ms: int, mult: float) -> None:
    jitter = random.uniform(0.8, 1.2)
    delay = base_ms * (mult ** max(0, attempt)) * jitter
    time.sleep(delay / 1000.0)


def normalize_lang(lang: str) -> str:
    l = str(lang or "").strip().lower()
    mapping = {
        "zh": "cn",
        "zh-cn": "cn",
        "zh_cn": "cn",
        "en-us": "en",
        "en_us": "en",
        "jp": "ja",
    }
    return mapping.get(l, l)


def call_xfyun_translate(provider: Provider, text: str, source_lang: str, target_lang: str) -> Tuple[str, Dict[str, Any]]:
    endpoint = provider.base_url
    u = urllib.parse.urlparse(endpoint)
    host = u.netloc
    path = u.path or "/v1/its"

    now = time.gmtime()
    date_str = time.strftime("%a, %d %b %Y %H:%M:%S GMT", now)

    request_line = f"POST {path} HTTP/1.1"
    signature_origin = f"host: {host}\ndate: {date_str}\n{request_line}"
    signature = base64.b64encode(
        hmac.new(provider.api_secret.encode("utf-8"), signature_origin.encode("utf-8"), hashlib.sha256).digest()
    ).decode("utf-8")

    authorization_origin = (
        f'api_key="{provider.api_key}", algorithm="hmac-sha256", '
        f'headers="host date request-line", signature="{signature}"'
    )
    authorization = base64.b64encode(authorization_origin.encode("utf-8")).decode("utf-8")

    query = urllib.parse.urlencode({
        "authorization": authorization,
        "host": host,
        "date": date_str,
    })
    url = f"{endpoint}?{query}"

    from_lang = normalize_lang(source_lang)
    to_lang = normalize_lang(target_lang)
    body = {
        "header": {"app_id": provider.app_id, "status": 3},
        "parameter": {"its": {"from": from_lang, "to": to_lang, "result": {}}},
        "payload": {
            "input_data": {
                "encoding": "utf8",
                "status": 3,
                "text": base64.b64encode(str(text).encode("utf-8")).decode("utf-8"),
            }
        },
    }
    payload = json.dumps(body).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=payload,
        headers={"Accept": "application/json", "Content-Type": "application/json"},
        method="POST",
    )

    with urllib.request.urlopen(req, timeout=(provider.timeout_ms / 1000.0)) as resp:
        raw = resp.read().decode("utf-8", errors="replace")
        data = json.loads(raw) if raw else {}

    header = data.get("header", {})
    code = int(header.get("code", -1))
    if code != 0:
        raise RuntimeError(f"XFYUN API error: code={code}, message={header.get('message')}, sid={header.get('sid')}")

    text_b64 = (((data.get("payload") or {}).get("result") or {}).get("text") or "")
    if not text_b64:
        raise RuntimeError("XFYUN API returned empty translation text")
    decoded = base64.b64decode(text_b64).decode("utf-8", errors="replace")
    decoded_json = json.loads(decoded)
    translated = (((decoded_json.get("trans_result") or {}).get("dst")) or "").strip()
    return translated, {"provider": provider.id, "model": provider.model or "xfyun-its"}


def call_openai_compatible(provider: Provider, messages: List[Dict[str, str]]) -> Tuple[str, Dict[str, Any]]:
    url = provider.base_url.rstrip("/") + "/chat/completions"
    headers: Dict[str, str] = {"Accept": "application/json", "Content-Type": "application/json"}
    if provider.requires_api_key and provider.api_key:
        headers["Authorization"] = f"Bearer {provider.api_key}"

    body = {
        "model": provider.model,
        "temperature": provider.temperature if provider.temperature is not None else 0,
        "top_p": provider.top_p if provider.top_p is not None else 0.1,
        "messages": messages,
    }

    payload = json.dumps(body).encode("utf-8")
    req = urllib.request.Request(url, data=payload, headers=headers, method="POST")
    timeout = provider.timeout_ms / 1000.0

    max_retries, backoff_ms, backoff_mult = _retry_params()
    last_err: Optional[Exception] = None
    for attempt in range(max_retries + 1):
        try:
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                raw = resp.read().decode("utf-8", errors="replace")
                data = json.loads(raw) if raw else {}
            last_err = None
            break
        except urllib.error.HTTPError as e:
            raw = ""
            try:
                raw = e.read().decode("utf-8", errors="replace")
            except Exception:
                pass
            code = getattr(e, "code", None)
            last_err = RuntimeError(f"LLM request failed: {code} {raw[:500]}")
            if code in (429, 500, 502, 503, 504) and attempt < max_retries:
                _sleep_backoff(attempt, backoff_ms, backoff_mult)
                continue
            raise last_err
        except urllib.error.URLError as e:
            last_err = RuntimeError(f"LLM request failed: {str(e)[:500]}")
            if attempt < max_retries:
                _sleep_backoff(attempt, backoff_ms, backoff_mult)
                continue
            raise last_err
    if last_err is not None:
        raise last_err

    content = (((data.get("choices") or [{}])[0].get("message") or {}).get("content") or "")
    usage = data.get("usage") or None
    model = data.get("model") or provider.model
    return str(content), {"usage": usage, "model": model, "provider": provider.id}


def _cache_key_text(text: str) -> str:
    if str(os.environ.get("TRANSLATE_NORMALIZE_WHITESPACE", "")).strip().lower() in ("1", "true", "yes", "on"):
        return normalize_text_whitespace(text)
    return text


def translate_text_with_cache(
    cache: SqliteCache,
    provider: Provider,
    lang_in: str,
    lang_out: str,
    glossary_hash: str,
    prompt_key: str,
    original_text: str,
    system_prompt: Optional[str] = None,
) -> Tuple[str, Dict[str, Any], bool]:
    t0 = time.perf_counter()
    # Dry-run mode: no external LLM call (for local smoke tests).
    dry_run = str(os.environ.get("TRANSLATE_DRY_RUN", "")).strip().lower() in ("1", "true", "yes", "on")

    # Build stable cache key: provider+model+langpair+glossary+prompt+textHash
    cache_text = _cache_key_text(original_text)
    text_hash = sha256_text(cache_text)
    key = sha256_text(
        "|".join([provider.id, provider.model, lang_in, lang_out, glossary_hash, prompt_key, text_hash])
    )
    cached = cache.get(key)
    if cached is not None:
        _log_timing("cache_hit", t0)
        return cached, {"cached": True}, True

    if dry_run:
        translated = original_text
        cache.set(
            key,
            {
                "provider_id": provider.id,
                "model": provider.model,
                "lang_in": lang_in,
                "lang_out": lang_out,
                "glossary_hash": glossary_hash,
                "prompt_version": prompt_key,
                "original_text": original_text,
                "translated_text": translated,
            },
        )
        _log_timing("dry_run", t0)
        return translated, {"dryRun": True}, False

    t_llm = time.perf_counter()
    if provider.kind == "xfyun":
        content, meta = call_xfyun_translate(provider, original_text, lang_in, lang_out)
    else:
        messages = build_translation_messages(lang_in, lang_out, original_text, system_prompt=system_prompt)
        content, meta = call_openai_compatible(provider, messages)
    _log_timing("llm_call", t_llm)
    translated = str(content).strip()

    cache.set(
        key,
        {
            "provider_id": provider.id,
            "model": provider.model,
            "lang_in": lang_in,
            "lang_out": lang_out,
            "glossary_hash": glossary_hash,
            "prompt_version": prompt_key,
            "original_text": original_text,
            "translated_text": translated,
        },
    )
    _log_timing("chunk_total", t0)
    return translated, meta, False


def _merge_usage(u1: Optional[Dict], u2: Optional[Dict]) -> Dict[str, int]:
    """Merge two usage dicts (input_tokens/output_tokens or prompt_tokens/completion_tokens) for logging."""
    def _in(d: Dict) -> int:
        return int(d.get("input_tokens") or d.get("prompt_tokens") or 0)

    def _out(d: Dict) -> int:
        return int(d.get("output_tokens") or d.get("completion_tokens") or 0)

    a = u1 or {}
    b = u2 or {}
    return {
        "input_tokens": _in(a) + _in(b),
        "output_tokens": _out(a) + _out(b),
    }


def translate_plain_text(
    text: str,
    provider: Provider,
    cache: SqliteCache,
    lang_in: str,
    lang_out: str,
    entries: List[Tuple[str, str]],
    synonyms: List[Tuple[str, List[str]]],
    glossary_hash: str,
    prompt_key: str,
    max_chars: int,
    is_markdown: bool,
    glossary_options: Optional[Dict[str, Any]] = None,
    system_prompt: Optional[str] = None,
) -> Tuple[str, Dict[str, Any]]:
    t_total = time.perf_counter()
    token_map: Dict[str, str] = {}
    work = text
    if is_markdown:
        t_md = time.perf_counter()
        work, token_map = protect_markdown(work)
        _log_timing("markdown_protect", t_md)

    # Term placeholders (locked)
    t_terms = time.perf_counter()
    work, placeholder_map, _used_sources = make_term_placeholders(work, entries, glossary_options)
    _log_timing("term_placeholders", t_terms)

    t_split = time.perf_counter()
    # 机器翻译 API（如讯飞）：逐行翻译，MT 会重排句子，合并后换行易错位；LLM：按段落翻译
    is_mt = provider.kind != "openai_compatible"
    preserve_lines = is_mt
    if preserve_lines:
        lines = work.split("\n")
        # MT 默认逐行，不合并，精确保留行边界
        segments = []
        for i, ln in enumerate(lines):
            if i > 0:
                segments.append(("\n", True))
            segments.append((ln, False))
    else:
        segments = split_paragraphs_keep_separators(work)
        merge_min = clamp_int(os.environ.get("TRANSLATE_MERGE_MIN_CHARS"), 50, 2000, 200)
        segments = merge_short_segments(segments, min_chars=merge_min, max_chars=max_chars)
    _log_timing("split_paragraphs", t_split)
    meta = {"chunks": 0, "cachedChunks": 0, "usage": {"input_tokens": 0, "output_tokens": 0}}
    chunk_concurrency = clamp_int(os.environ.get("TRANSLATE_CHUNK_CONCURRENCY"), 1, 16, 4)
    # MT 默认 segment 并发 4，LLM 默认 1
    segment_concurrency = clamp_int(
        os.environ.get("TRANSLATE_SEGMENT_CONCURRENCY"), 1, 16,
        4 if is_mt else 1,
    )

    def translate_one_segment(seg: str) -> Tuple[str, int, int, Dict[str, int]]:
        """Translate one segment. Returns (joined_text, chunks, cached_chunks, usage)."""
        subchunks = split_long_text(seg, max_chars=max_chars)
        translated_sub: List[str] = []
        seg_chunks = 0
        seg_cached = 0
        seg_usage: Dict[str, int] = {}
        if chunk_concurrency <= 1 or len(subchunks) <= 1:
            for ch in subchunks:
                translated, _m, was_cached = translate_text_with_cache(
                    cache=cache,
                    provider=provider,
                    lang_in=lang_in,
                    lang_out=lang_out,
                    glossary_hash=glossary_hash,
                    prompt_key=prompt_key,
                    original_text=ch,
                    system_prompt=system_prompt,
                )
                translated_sub.append(translated)
                seg_chunks += 1
                if was_cached:
                    seg_cached += 1
                seg_usage = _merge_usage(seg_usage, _m.get("usage"))
        else:
            futures = []
            with ThreadPoolExecutor(max_workers=chunk_concurrency) as executor:
                for idx, ch in enumerate(subchunks):
                    seg_chunks += 1
                    futures.append(
                        (idx, executor.submit(
                            translate_text_with_cache,
                            cache,
                            provider,
                            lang_in,
                            lang_out,
                            glossary_hash,
                            prompt_key,
                            ch,
                            system_prompt,
                        ))
                    )
                results = [None] * len(subchunks)
                for idx, fut in futures:
                    translated, _m, was_cached = fut.result()
                    if was_cached:
                        seg_cached += 1
                    seg_usage = _merge_usage(seg_usage, _m.get("usage"))
                    results[idx] = translated
                translated_sub = [r or "" for r in results]
        joined = "".join(translated_sub)
        if preserve_lines:
            joined = _restore_line_merge_delim(joined)
        return joined, seg_chunks, seg_cached, seg_usage

    items: List[Tuple[bool, str]] = [(is_sep, seg) for seg, is_sep in segments]
    content_indices = [i for i, (is_sep, seg) in enumerate(items) if not is_sep and seg.strip()]

    if segment_concurrency <= 1 or len(content_indices) <= 1:
        translated_map: Dict[int, str] = {}
        for i in content_indices:
            seg = items[i][1]
            joined, sc, sc_cached, su = translate_one_segment(seg)
            translated_map[i] = joined
            meta["chunks"] += sc
            meta["cachedChunks"] += sc_cached
            meta["usage"] = _merge_usage(meta.get("usage"), su)
    else:
        translated_map = {}
        with ThreadPoolExecutor(max_workers=segment_concurrency) as executor:
            futures_map = {i: executor.submit(translate_one_segment, items[i][1]) for i in content_indices}
            for i, fut in futures_map.items():
                joined, sc, sc_cached, su = fut.result()
                translated_map[i] = joined
                meta["chunks"] += sc
                meta["cachedChunks"] += sc_cached
                meta["usage"] = _merge_usage(meta.get("usage"), su)

    out_parts: List[str] = []
    for i, (is_sep, seg) in enumerate(items):
        if is_sep:
            out_parts.append(seg)
        elif not seg.strip():
            out_parts.append(seg)
        else:
            out_parts.append(translated_map[i])

    out = "".join(out_parts)
    out = restore_placeholders(out, placeholder_map)
    out = apply_synonym_fixes(out, synonyms)
    if is_markdown:
        out = restore_markdown(out, token_map)
    _log_timing("translate_plain_total", t_total)
    return out, meta


def translate_json_value(
    value: Any,
    provider: Provider,
    cache: SqliteCache,
    lang_in: str,
    lang_out: str,
    entries: List[Tuple[str, str]],
    synonyms: List[Tuple[str, List[str]]],
    glossary_hash: str,
    prompt_key: str,
    max_chars: int,
    glossary_options: Optional[Dict[str, Any]] = None,
    system_prompt: Optional[str] = None,
) -> Tuple[Any, Dict[str, Any]]:
    meta = {"strings": 0, "chunks": 0, "cachedChunks": 0, "usage": {"input_tokens": 0, "output_tokens": 0}}

    def walk(v: Any) -> Any:
        if isinstance(v, str):
            meta["strings"] += 1
            translated, m = translate_plain_text(
                text=v,
                provider=provider,
                cache=cache,
                lang_in=lang_in,
                lang_out=lang_out,
                entries=entries,
                synonyms=synonyms,
                glossary_hash=glossary_hash,
                prompt_key=prompt_key,
                max_chars=max_chars,
                is_markdown=False,
                glossary_options=glossary_options,
                system_prompt=system_prompt,
            )
            meta["chunks"] += m.get("chunks", 0)
            meta["cachedChunks"] += m.get("cachedChunks", 0)
            meta["usage"] = _merge_usage(meta.get("usage"), m.get("usage"))
            return translated
        if isinstance(v, list):
            return [walk(x) for x in v]
        if isinstance(v, dict):
            return {k: walk(val) for k, val in v.items()}
        return v

    return walk(value), meta


def translate_xml_file(
    input_path: Path,
    output_path: Path,
    provider: Provider,
    cache: SqliteCache,
    lang_in: str,
    lang_out: str,
    entries: List[Tuple[str, str]],
    synonyms: List[Tuple[str, List[str]]],
    glossary_hash: str,
    prompt_key: str,
    max_chars: int,
    glossary_options: Optional[Dict[str, Any]] = None,
    system_prompt: Optional[str] = None,
) -> Dict[str, Any]:
    try:
        tree = ET.parse(str(input_path))
    except Exception as e:
        raise RuntimeError(f"failed to parse xml: {e}") from e

    meta = {"strings": 0, "chunks": 0, "cachedChunks": 0, "usage": {"input_tokens": 0, "output_tokens": 0}}

    def _translate_text(text: Optional[str]) -> Optional[str]:
        if text is None:
            return None
        if not isinstance(text, str):
            text = str(text)
        if not text.strip():
            return text

        # Preserve surrounding whitespace/indentation while translating only the core text.
        lead_m = re.match(r"^\s*", text)
        tail_m = re.match(r"\s*$", text)
        lead = lead_m.group(0) if lead_m else ""
        tail = tail_m.group(0) if tail_m else ""
        core = text[len(lead):]
        if tail:
            core = core[: -len(tail)]

        meta["strings"] += 1
        translated, m = translate_plain_text(
            text=core,
            provider=provider,
            cache=cache,
            lang_in=lang_in,
            lang_out=lang_out,
            entries=entries,
            synonyms=synonyms,
            glossary_hash=glossary_hash,
            prompt_key=prompt_key,
            max_chars=max_chars,
            is_markdown=False,
            glossary_options=glossary_options,
            system_prompt=system_prompt,
        )
        meta["chunks"] += m.get("chunks", 0)
        meta["cachedChunks"] += m.get("cachedChunks", 0)
        meta["usage"] = _merge_usage(meta.get("usage"), m.get("usage"))
        return f"{lead}{translated}{tail}"

    def _walk(elem: ET.Element) -> None:
        elem.text = _translate_text(elem.text)
        for child in list(elem):
            _walk(child)
            child.tail = _translate_text(child.tail)

    _walk(tree.getroot())
    ensure_dir(output_path.parent)
    tree.write(str(output_path), encoding="utf-8", xml_declaration=True)
    return meta


def iter_docx_text_targets(doc):
    # paragraphs
    for p in doc.paragraphs:
        yield p
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    # headers/footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for p in section.footer.paragraphs:
            yield p
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def translate_docx_file(
    input_path: Path,
    output_path: Path,
    provider: Provider,
    cache: SqliteCache,
    lang_in: str,
    lang_out: str,
    entries: List[Tuple[str, str]],
    synonyms: List[Tuple[str, List[str]]],
    glossary_hash: str,
    prompt_key: str,
    max_chars: int,
    glossary_options: Optional[Dict[str, Any]] = None,
    system_prompt: Optional[str] = None,
) -> Dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise RuntimeError(
            "python-docx is required to translate .docx files. Please install it in your Python environment."
        ) from e

    doc = Document(str(input_path))
    meta = {"paragraphs": 0, "chunks": 0, "cachedChunks": 0, "usage": {"input_tokens": 0, "output_tokens": 0}}

    def _build_run_marked_text(runs) -> Tuple[str, List[int]]:
        parts: List[str] = []
        run_indexes: List[int] = []
        idx = 1
        for i, r in enumerate(runs):
            txt = r.text or ""
            if not txt:
                continue
            token = f"[[RUN_{idx:04d}]]"
            parts.append(token + txt)
            run_indexes.append(i)
            idx += 1
        return "".join(parts), run_indexes

    def _split_by_run_markers(text: str) -> Optional[List[str]]:
        # Expect pattern: [[RUN_0001]]...[[RUN_0002]]...
        parts = re.split(r"\[\[RUN_\d{4}\]\]", text)
        # parts[0] is preamble before first marker (should be empty/whitespace)
        if len(parts) <= 1:
            return None
        if parts[0].strip():
            # Unexpected content before first marker
            return None
        return parts[1:]

    for p in iter_docx_text_targets(doc):
        txt = p.text
        if not txt or not txt.strip():
            continue
        meta["paragraphs"] += 1
        runs = list(p.runs)
        if len(runs) <= 1:
            translated, m = translate_plain_text(
                text=txt,
                provider=provider,
                cache=cache,
                lang_in=lang_in,
                lang_out=lang_out,
                entries=entries,
                synonyms=synonyms,
                glossary_hash=glossary_hash,
                prompt_key=prompt_key,
                max_chars=max_chars,
                is_markdown=False,
                glossary_options=glossary_options,
                system_prompt=system_prompt,
            )
            meta["chunks"] += m.get("chunks", 0)
            meta["cachedChunks"] += m.get("cachedChunks", 0)
            meta["usage"] = _merge_usage(meta.get("usage"), m.get("usage"))
            p.text = translated
            continue

        marked_text, run_indexes = _build_run_marked_text(runs)
        if not marked_text:
            continue
        run_prompt = system_prompt or _default_system_prompt(lang_in, lang_out)
        run_prompt += "\n- Preserve tokens like [[RUN_0001]] exactly and do not remove them.\n"
        translated, m = translate_plain_text(
            text=marked_text,
            provider=provider,
            cache=cache,
            lang_in=lang_in,
            lang_out=lang_out,
            entries=entries,
            synonyms=synonyms,
            glossary_hash=glossary_hash,
            prompt_key=prompt_key,
            max_chars=max_chars,
            is_markdown=False,
            glossary_options=glossary_options,
            system_prompt=run_prompt,
        )
        meta["chunks"] += m.get("chunks", 0)
        meta["cachedChunks"] += m.get("cachedChunks", 0)
        meta["usage"] = _merge_usage(meta.get("usage"), m.get("usage"))

        segments = _split_by_run_markers(translated)
        if segments is None or len(segments) != len(run_indexes):
            if LOG_TIMING:
                sys.stderr.write(
                    f"[translate-kernel] run_marker_mismatch: expected={len(run_indexes)} got={0 if segments is None else len(segments)}\n"
                )
            # Fallback to paragraph-level translation to avoid losing output
            p.text = translated
            continue

        # Assign translated text back to runs, preserving run-level styles
        for seg, run_idx in zip(segments, run_indexes):
            runs[run_idx].text = seg
        # Clear text of runs that were not mapped to avoid duplicate content
        for i, r in enumerate(runs):
            if i not in run_indexes:
                r.text = ""

    ensure_dir(output_path.parent)
    doc.save(str(output_path))
    return meta


def translate_file(
    input_path: Path,
    output_path: Path,
    provider: Provider,
    cache: SqliteCache,
    lang_in: str,
    lang_out: str,
    entries: List[Tuple[str, str]],
    synonyms: List[Tuple[str, List[str]]],
    glossary_hash: str,
    prompt_key: str,
    max_chars: int,
    glossary_options: Optional[Dict[str, Any]] = None,
    system_prompt: Optional[str] = None,
) -> Dict[str, Any]:
    ext = input_path.suffix.lower()
    if ext == ".pdf":
        from pdf_translate import translate_pdf_file
        def _pdf_translate(text: str) -> str:
            pdf_prompt = (system_prompt or _default_system_prompt(lang_in, lang_out)) + "\n- Preserve placeholders like {v0} exactly.\n"
            translated, m = translate_plain_text(
                text=text,
                provider=provider,
                cache=cache,
                lang_in=lang_in,
                lang_out=lang_out,
                entries=entries,
                synonyms=synonyms,
                glossary_hash=glossary_hash,
                prompt_key=prompt_key,
                max_chars=max_chars,
                is_markdown=False,
                glossary_options=glossary_options,
                system_prompt=pdf_prompt,
            )
            return translated

        thread = clamp_int(os.environ.get("TRANSLATE_PDF_THREAD"), 1, 16, 2)
        return translate_pdf_file(
            input_path=input_path,
            output_path=output_path,
            translate_text_fn=_pdf_translate,
            lang_in=lang_in,
            lang_out=lang_out,
            thread=thread,
        )
    if ext == ".docx":
        return translate_docx_file(
            input_path=input_path,
            output_path=output_path,
            provider=provider,
            cache=cache,
            lang_in=lang_in,
            lang_out=lang_out,
            entries=entries,
            synonyms=synonyms,
            glossary_hash=glossary_hash,
            max_chars=max_chars,
        )
    if ext in (".txt", ".md"):
        raw = input_path.read_text(encoding="utf-8", errors="ignore")
        translated, meta = translate_plain_text(
            text=raw,
            provider=provider,
            cache=cache,
            lang_in=lang_in,
            lang_out=lang_out,
            entries=entries,
            synonyms=synonyms,
            glossary_hash=glossary_hash,
            prompt_key=prompt_key,
            max_chars=max_chars,
            is_markdown=(ext == ".md"),
            glossary_options=glossary_options,
            system_prompt=system_prompt,
        )
        ensure_dir(output_path.parent)
        output_path.write_text(translated, encoding="utf-8")
        return meta
    if ext == ".json":
        raw = input_path.read_text(encoding="utf-8", errors="ignore")
        obj = json.loads(raw)
        translated_obj, meta = translate_json_value(
            value=obj,
            provider=provider,
            cache=cache,
            lang_in=lang_in,
            lang_out=lang_out,
            entries=entries,
            synonyms=synonyms,
            glossary_hash=glossary_hash,
            prompt_key=prompt_key,
            max_chars=max_chars,
            glossary_options=glossary_options,
            system_prompt=system_prompt,
        )
        ensure_dir(output_path.parent)
        output_path.write_text(json.dumps(translated_obj, ensure_ascii=False, indent=2), encoding="utf-8")
        return meta
    if ext == ".xml":
        return translate_xml_file(
            input_path=input_path,
            output_path=output_path,
            provider=provider,
            cache=cache,
            lang_in=lang_in,
            lang_out=lang_out,
            entries=entries,
            synonyms=synonyms,
            glossary_hash=glossary_hash,
            prompt_key=prompt_key,
            max_chars=max_chars,
            glossary_options=glossary_options,
            system_prompt=system_prompt,
        )
    raise RuntimeError(f"unsupported file type: {ext}")


def _load_prompt_config(prompts_file: Path, prompt_key: str) -> Optional[str]:
    raw = safe_json_load(prompts_file)
    if not isinstance(raw, dict):
        return None
    entry = raw.get(prompt_key)
    if not isinstance(entry, dict):
        return None
    system_lines = entry.get("system")
    if not isinstance(system_lines, list):
        return None
    system = "\n".join(str(x) for x in system_lines if str(x).strip())
    return system or None


def _resolve_prompt(
    source_lang: str,
    target_lang: str,
    prompts_file: Optional[Path],
    prompt_key: str,
) -> Tuple[Optional[str], str]:
    system_prompt = None
    if prompts_file and prompts_file.exists():
        system_prompt = _load_prompt_config(prompts_file, prompt_key)
    if system_prompt:
        system_prompt = (
            system_prompt.replace("{{source_lang}}", source_lang)
            .replace("{{target_lang}}", target_lang)
            .replace("{source_lang}", source_lang)
            .replace("{target_lang}", target_lang)
        )
    else:
        system_prompt = _default_system_prompt(source_lang, target_lang)
    prompt_hash = sha256_text(system_prompt or "")
    prompt_key = f"{PROMPT_VERSION}:{prompt_hash[:10]}"
    return system_prompt, prompt_key


def cmd_translate_file(args: argparse.Namespace) -> int:
    t_all = time.perf_counter()
    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()
    providers_file = Path(args.providers_file).resolve()
    glossary_db = Path(args.glossary_db).resolve() if args.glossary_db else None
    glossary_domain = str(args.glossary_domain or "").strip() or None
    max_chars = int(args.max_chars)

    providers = load_providers(providers_file)
    provider = resolve_provider(providers, args.provider_id)
    ok, reason = llm_available(provider)
    if not ok:
        raise RuntimeError(f"LLM not available: {reason}")

    glossary_disabled = str(os.environ.get("TRANSLATE_GLOSSARY_DISABLED", "")).strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )
    t_glossary = time.perf_counter()
    if glossary_disabled:
        entries, synonyms, glossary_hash, glossary_options = [], [], "disabled", {}
    elif not glossary_db:
        raise RuntimeError("Missing --glossary-db (SQLite glossary is required)")
    else:
        entries, synonyms, glossary_hash, glossary_options = glossary_load_from_sqlite(
            glossary_db=glossary_db,
            source_lang=args.source_lang,
            target_lang=args.target_lang,
            domain=glossary_domain,
        )
    _log_timing("glossary_load", t_glossary)

    prompts_file = None
    prompts_file_raw = str(os.environ.get("TRANSLATE_PROMPTS_FILE", "") or "").strip()
    if prompts_file_raw:
        prompts_file = Path(prompts_file_raw).expanduser()
    else:
        prompts_file = Path(__file__).resolve().parents[1] / "src" / "config" / "smartSearchPrompts.json"

    prompt_key_env = str(os.environ.get("TRANSLATE_PROMPT_KEY", "") or "").strip() or "translationKernel"
    t_prompt = time.perf_counter()
    system_prompt, prompt_key = _resolve_prompt(
        source_lang=args.source_lang,
        target_lang=args.target_lang,
        prompts_file=prompts_file,
        prompt_key=prompt_key_env,
    )
    _log_timing("prompt_resolve", t_prompt)

    # Use the same SQLite file for glossary and cache (different tables).
    raw_cache_db = str(os.environ.get("TRANSLATE_CACHE_DB", "") or "").strip()
    if raw_cache_db:
        cache_db = Path(raw_cache_db).expanduser()
    else:
        cache_db = glossary_db
    t_cache = time.perf_counter()
    cache = SqliteCache(cache_db)
    _log_timing("cache_init", t_cache)

    t_translate = time.perf_counter()
    meta = translate_file(
        input_path=input_path,
        output_path=output_path,
        provider=provider,
        cache=cache,
        lang_in=args.source_lang,
        lang_out=args.target_lang,
        entries=entries,
        synonyms=synonyms,
        glossary_hash=glossary_hash,
        prompt_key=prompt_key,
        max_chars=max_chars,
        glossary_options=glossary_options,
        system_prompt=system_prompt,
    )
    _log_timing("translate_file", t_translate)

    out = {
        "ok": True,
        "meta": {
            **meta,
            "provider": {"id": provider.id, "model": provider.model},
            "promptVersion": prompt_key,
            "glossaryHash": glossary_hash,
        },
        "warnings": [
            "docx translation may lose run-level styling (bold/italic) in this first version."
        ]
        if input_path.suffix.lower() == ".docx"
        else [],
    }
    sys.stdout.write(json.dumps(out, ensure_ascii=False))
    _log_timing("translate_total", t_all)
    return 0


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(prog="translation-kernel")
    sub = p.add_subparsers(dest="command", required=True)

    t = sub.add_parser("translate-file", help="Translate a file and write output")
    t.add_argument("--input", required=True)
    t.add_argument("--output", required=True)
    t.add_argument("--source-lang", required=True)
    t.add_argument("--target-lang", required=True)
    t.add_argument("--providers-file", required=True)
    t.add_argument("--provider-id", required=False, default=None)
    t.add_argument("--glossary-db", required=True)
    t.add_argument("--glossary-domain", required=False, default=None)
    t.add_argument("--max-chars", required=False, default=str(int(os.environ.get("TRANSLATE_CHUNK_MAX_CHARS", "1200"))))
    t.set_defaults(func=cmd_translate_file)

    return p


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    try:
        return int(args.func(args))
    except Exception as e:
        err = {"ok": False, "error": str(e), "type": e.__class__.__name__}
        sys.stderr.write(json.dumps(err, ensure_ascii=False) + "\n")
        return 2


if __name__ == "__main__":
    raise SystemExit(main())

